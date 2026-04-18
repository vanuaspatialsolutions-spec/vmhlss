"""
Document Intelligence Service

Processes uploaded documents (PDF, DOCX, XLSX, images) to extract geospatial,
hazard, soil, engineering, and policy information using Claude API.
"""

import json
import logging
import time
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import base64

import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import pypdf
from docx import Document
from openpyxl import load_workbook
import anthropic
from sqlalchemy.orm import Session

from app.models.knowledge_base import KnowledgeBaseRecord, InformationType
from app.models.gazetteer import VanuatuPlace
from app.database import get_db

logger = logging.getLogger(__name__)

EXTRACTION_SYSTEM_PROMPT = """
You are a geospatial data extraction specialist working on land use planning for Vanuatu.
Extract all spatially relevant information from the document text provided.

Return ONLY a valid JSON object with this exact structure:
{
  "locations_events": [
    {
      "description": "text describing the location or event",
      "place_name": "name if mentioned",
      "coordinates": {"lat": null, "lon": null},
      "event_type": "landslide|flood|cyclone|volcanic|earthquake|tsunami|other",
      "date": "YYYY-MM-DD or null",
      "intensity": "description of intensity if mentioned",
      "confidence": 0.0-1.0
    }
  ],
  "hazard_data": [
    {
      "description": "extracted hazard description",
      "hazard_type": "cyclone|tsunami|volcanic|flood|earthquake|landslide",
      "location_description": "where this applies",
      "buffer_distance_m": null,
      "exclusion_zone": false,
      "confidence": 0.0-1.0
    }
  ],
  "soil_agricultural": [
    {
      "description": "soil or agricultural observation",
      "soil_type": "text or null",
      "fao_class": "S1|S2|S3|S4|S5|N1|N2 or null",
      "suitable_crops": [],
      "drainage_class": "text or null",
      "location_description": "where this applies",
      "confidence": 0.0-1.0
    }
  ],
  "engineering_data": [
    {
      "description": "engineering observation or recommendation",
      "parameter_type": "foundation_depth|bearing_capacity|slope_stability|seismic_design|other",
      "value": "extracted value with units",
      "location_description": "where this applies",
      "confidence": 0.0-1.0
    }
  ],
  "policy_legal": [
    {
      "description": "policy or legal reference",
      "type": "exclusion|buffer|condition|recommendation|directive",
      "buffer_distance_m": null,
      "location_description": "where this applies",
      "legal_authority": "issuing authority if mentioned",
      "confidence": 0.0-1.0
    }
  ],
  "document_summary": "two sentence summary of the document",
  "document_type": "hazard_assessment|agricultural_report|engineering_study|eia|policy|other"
}
"""


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF. Try pypdf first; fall back to OCR if text is sparse.
    """
    logger.info(f"Extracting text from PDF: {file_path}")
    text = ""

    try:
        with open(file_path, "rb") as pdf_file:
            reader = pypdf.PdfReader(pdf_file)
            for page in reader.pages:
                text += page.extract_text() + "\n"

        if len(text.strip()) > 100:
            logger.info(f"Extracted {len(text)} characters using pypdf")
            return text
    except Exception as e:
        logger.warning(f"pypdf extraction failed: {e}")

    logger.info("Text extraction sparse; falling back to Tesseract OCR")
    try:
        images = convert_from_path(file_path)
        for img in images:
            text += pytesseract.image_to_string(img, lang='eng') + "\n"
        logger.info(f"OCR extracted {len(text)} characters")
        return text
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        raise


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX using python-docx.
    """
    logger.info(f"Extracting text from DOCX: {file_path}")
    doc = Document(file_path)
    text = ""

    for para in doc.paragraphs:
        text += para.text + "\n"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " | "
            text += "\n"

    logger.info(f"Extracted {len(text)} characters from DOCX")
    return text


def extract_text_from_xlsx(file_path: str) -> str:
    """
    Extract text from XLSX using openpyxl.
    """
    logger.info(f"Extracting text from XLSX: {file_path}")
    wb = load_workbook(file_path)
    text = ""

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text += f"\n=== Sheet: {sheet_name} ===\n"

        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell else "" for cell in row)
            text += row_text + "\n"

    logger.info(f"Extracted {len(text)} characters from XLSX")
    return text


def extract_text_from_image(file_path: str) -> str:
    """
    Extract text from image using Tesseract OCR.
    """
    logger.info(f"Extracting text from image: {file_path}")
    img = Image.open(file_path)
    text = pytesseract.image_to_string(img, lang='eng')
    logger.info(f"Extracted {len(text)} characters from image")
    return text


def call_claude_extraction(text: str, document_name: str, max_retries: int = 3) -> Dict:
    """
    Call Anthropic Claude API to extract structured data from document text.
    Implements exponential backoff for rate limit handling.
    """
    logger.info(f"Calling Claude API for extraction: {document_name}")

    client = anthropic.Anthropic()

    for attempt in range(max_retries):
        try:
            message = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                system=EXTRACTION_SYSTEM_PROMPT,
                messages=[
                    {
                        "role": "user",
                        "content": f"Extract geospatial information from this document:\n\n{text[:8000]}"
                    }
                ]
            )

            response_text = message.content[0].text

            try:
                result = json.loads(response_text)
                logger.info(f"Extraction successful: {document_name}")
                return result
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse Claude response as JSON: {e}")
                logger.error(f"Response: {response_text[:500]}")
                raise

        except anthropic.RateLimitError as e:
            if attempt < max_retries - 1:
                wait_time = 2 ** attempt
                logger.warning(f"Rate limited. Retrying in {wait_time}s (attempt {attempt+1}/{max_retries})")
                time.sleep(wait_time)
            else:
                logger.error(f"Max retries exceeded for rate limit")
                raise
        except anthropic.APIError as e:
            logger.error(f"Claude API error: {e}")
            raise

    raise RuntimeError(f"Failed to extract data from {document_name} after {max_retries} attempts")


def parse_coordinate_text(text: str) -> Optional[Tuple[float, float]]:
    """
    Parse coordinate text in multiple formats.
    Returns (lat, lon) tuple or None.

    Supported formats:
    - Decimal degrees: "-15.376, 166.959"
    - DMS: "15°22'33\"S 166°57'32\"E"
    """
    text = text.strip()

    if not text:
        return None

    # Try decimal degrees
    try:
        parts = [p.strip() for p in text.replace(',', ' ').split()]
        if len(parts) >= 2:
            lat = float(parts[0])
            lon = float(parts[1])
            if -90 <= lat <= 90 and -180 <= lon <= 180:
                return (lat, lon)
    except (ValueError, IndexError):
        pass

    # Try DMS format (simplified)
    try:
        import re
        dms_pattern = r"(\d+)°(\d+)'([\d.]+)\"([NSEW])\s+(\d+)°(\d+)'([\d.]+)\"([NSEW])"
        match = re.search(dms_pattern, text)
        if match:
            lat_d, lat_m, lat_s, lat_dir, lon_d, lon_m, lon_s, lon_dir = match.groups()

            lat = float(lat_d) + float(lat_m)/60 + float(lat_s)/3600
            lon = float(lon_d) + float(lon_m)/60 + float(lon_s)/3600

            if lat_dir in ['S', 's']:
                lat = -lat
            if lon_dir in ['W', 'w']:
                lon = -lon

            return (lat, lon)
    except Exception:
        pass

    return None


def resolve_place_names(items: List[Dict], db: Session) -> List[Dict]:
    """
    Resolve place names using Vanuatu gazetteer.
    Updates location_source field and coordinates.
    """
    logger.info(f"Resolving place names for {len(items)} items")

    for item in items:
        if not item.get('place_name'):
            item['location_source'] = 'unresolved'
            continue

        place_name = item['place_name']

        # Try gazetteer lookup with ILIKE
        place = db.query(VanuatuPlace).filter(
            VanuatuPlace.name.ilike(f"%{place_name}%")
        ).first()

        if place:
            item['coordinates'] = {
                'lat': float(place.latitude),
                'lon': float(place.longitude)
            }
            item['location_source'] = 'gazetteer'
            logger.debug(f"Resolved {place_name} to {place.latitude}, {place.longitude}")
        elif item.get('coordinates', {}).get('lat') is not None:
            item['location_source'] = 'coordinates'
            logger.debug(f"Using provided coordinates for {place_name}")
        else:
            item['location_source'] = 'unresolved'
            logger.warning(f"Could not resolve place name: {place_name}")

    return items


def process_document(
    file_path: str,
    document_name: str,
    uploaded_by: str,
    db: Session
) -> Dict:
    """
    Main document processing pipeline.

    1. Detect file type
    2. Extract text
    3. Call Claude API
    4. Resolve place names
    5. Store in knowledge base
    6. Return results
    """
    logger.info(f"Processing document: {document_name}")

    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    # Step 1: Extract text based on file type
    if suffix == '.pdf':
        text = extract_text_from_pdf(str(file_path))
    elif suffix == '.docx':
        text = extract_text_from_docx(str(file_path))
    elif suffix == '.xlsx':
        text = extract_text_from_xlsx(str(file_path))
    elif suffix in ['.png', '.jpg', '.jpeg', '.tiff']:
        text = extract_text_from_image(str(file_path))
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    if not text or len(text.strip()) < 50:
        raise ValueError("Document text extraction yielded insufficient content")

    # Step 2: Call Claude API
    extraction_result = call_claude_extraction(text, document_name)

    # Step 3: Resolve place names
    if extraction_result.get('locations_events'):
        extraction_result['locations_events'] = resolve_place_names(
            extraction_result['locations_events'],
            db
        )

    # Step 4: Store in knowledge base
    created_record_ids = store_extracted_items(
        extraction_result,
        document_name,
        str(file_path),
        uploaded_by,
        db
    )

    logger.info(f"Document processing complete: {len(created_record_ids)} records created")

    return {
        'document_name': document_name,
        'document_type': extraction_result.get('document_type'),
        'document_summary': extraction_result.get('document_summary'),
        'records_created': len(created_record_ids),
        'record_ids': created_record_ids,
        'extraction_details': extraction_result
    }


def store_extracted_items(
    items: Dict,
    document_name: str,
    document_path: str,
    extracted_by: str,
    db: Session
) -> List[str]:
    """
    Store extracted items as knowledge base records.
    Each item becomes a separate record with status='pending'.
    """
    logger.info(f"Storing extracted items from {document_name}")

    record_ids = []

    # Process locations_events
    for item in items.get('locations_events', []):
        record = KnowledgeBaseRecord(
            information_type=InformationType.HAZARD_EVENT,
            extracted_text=item.get('description', ''),
            interpreted_value=json.dumps(item),
            place_name=item.get('place_name'),
            coordinates=item.get('coordinates'),
            confidence_score=item.get('confidence', 0.5),
            location_source=item.get('location_source', 'unresolved'),
            source_document=document_name,
            source_document_path=document_path,
            extracted_by=extracted_by,
            status='pending'
        )
        db.add(record)
        db.flush()
        record_ids.append(str(record.id))

    # Process hazard_data
    for item in items.get('hazard_data', []):
        record = KnowledgeBaseRecord(
            information_type=InformationType.HAZARD_ZONE,
            extracted_text=item.get('description', ''),
            interpreted_value=json.dumps(item),
            place_name=item.get('location_description'),
            confidence_score=item.get('confidence', 0.5),
            location_source='document',
            source_document=document_name,
            source_document_path=document_path,
            extracted_by=extracted_by,
            status='pending'
        )
        db.add(record)
        db.flush()
        record_ids.append(str(record.id))

    # Process soil_agricultural
    for item in items.get('soil_agricultural', []):
        record = KnowledgeBaseRecord(
            information_type=InformationType.SOIL_DATA,
            extracted_text=item.get('description', ''),
            interpreted_value=json.dumps(item),
            place_name=item.get('location_description'),
            confidence_score=item.get('confidence', 0.5),
            location_source='document',
            source_document=document_name,
            source_document_path=document_path,
            extracted_by=extracted_by,
            status='pending'
        )
        db.add(record)
        db.flush()
        record_ids.append(str(record.id))

    # Process engineering_data
    for item in items.get('engineering_data', []):
        record = KnowledgeBaseRecord(
            information_type=InformationType.ENGINEERING_DATA,
            extracted_text=item.get('description', ''),
            interpreted_value=json.dumps(item),
            place_name=item.get('location_description'),
            confidence_score=item.get('confidence', 0.5),
            location_source='document',
            source_document=document_name,
            source_document_path=document_path,
            extracted_by=extracted_by,
            status='pending'
        )
        db.add(record)
        db.flush()
        record_ids.append(str(record.id))

    # Process policy_legal
    for item in items.get('policy_legal', []):
        record = KnowledgeBaseRecord(
            information_type=InformationType.POLICY_LEGAL,
            extracted_text=item.get('description', ''),
            interpreted_value=json.dumps(item),
            place_name=item.get('location_description'),
            confidence_score=item.get('confidence', 0.5),
            location_source='document',
            source_document=document_name,
            source_document_path=document_path,
            extracted_by=extracted_by,
            status='pending'
        )
        db.add(record)
        db.flush()
        record_ids.append(str(record.id))

    db.commit()
    logger.info(f"Stored {len(record_ids)} knowledge base records")

    return record_ids
